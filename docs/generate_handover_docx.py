"""
Generate the handover report as a Word .docx file.

Same content as docs/cyberscore-final-report.pdf but in a format that opens
natively in Word, Google Docs, and Pages so the supervisor (or anyone)
can edit, comment, and re-export to PDF without needing the Python script.

Run with:
    python3 docs/generate_handover_docx.py

Writes to: docs/cyberscore-final-report.docx

Style is plain English, slightly formal, written in Saanvi's first person
voice. No em-dashes, no decorative emojis, no status emojis.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/saanvivishal/Desktop/cyberscore/docs/cyberscore-final-report.docx"

# Colors (RGB)
BRAND_BLUE = RGBColor(0x3B, 0x82, 0xF6)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _add_page_number(footer_para):
    """Add a page-number field to a footer paragraph."""
    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _set_cell_shading(cell, color_hex):
    """Apply background colour to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    """Add a heading with brand colour."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(28)
        run.font.color.rgb = INK
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = BRAND_BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = INK
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = INK
        run.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)


def add_para(doc, text, size=11, bold=False, italic=False, color=INK, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullets(doc, items):
    """Add a bullet list. Items are strings."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, col_widths_inches=None):
    """Add a table with a brand-blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Set column widths if provided
    if col_widths_inches:
        for i, width in enumerate(col_widths_inches):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        _set_cell_shading(hdr_cells[i], "3B82F6")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            # Light alternating shade
            if r_idx % 2 == 1:
                _set_cell_shading(cells[c_idx], "F1F5F9")

    # Spacing after the table
    doc.add_paragraph()


# ============================================================
# Build the document
# ============================================================

doc = Document()

# Page setup: A4-ish margins
for section in doc.sections:
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Footer with page number
for section in doc.sections:
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "CyberScore. Final report. Saanvi Vishal, IIIT Bangalore. Page "
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED
    _add_page_number(fp)


# ----- COVER -----
for _ in range(4):
    doc.add_paragraph()

add_heading(doc, "CyberScore", level=0)
add_para(
    doc,
    "A Mobile-First Cybersecurity Health Scorecard for Organisations",
    size=14,
    italic=True,
    color=MUTED,
    align="center",
)
doc.add_paragraph()
add_para(doc, "Final Project Report", size=16, bold=True, align="center")
doc.add_paragraph()
doc.add_paragraph()

meta_lines = [
    "Submitted by: Saanvi Vishal (IMT2021043)",
    "Mentor: Mohan Ram C, FISST",
    "Institution: International Institute of Information Technology, Bangalore",
    "Project duration: 15 April 2026 to 19 May 2026",
    "Repository: https://github.com/saanvivishal/cyberscore",
    "Live API: https://cyberscore-api.vercel.app",
    "Demo video: hosted on IIIT Bangalore SharePoint (link in handover email)",
]
for line in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = INK

doc.add_page_break()


# ----- 1. Abstract -----
add_heading(doc, "1. Abstract")
add_para(
    doc,
    "CyberScore is a mobile-first SaaS that lets a company check its own cybersecurity health in about thirty minutes. The user answers 46 questions split across three areas: People (workforce hygiene), Process (operational discipline), and Company (governance and risk). The app then gives them a live numeric score from zero to one hundred, shows which areas are weak, and offers personalised remediation suggestions. There is also a built-in AI advisor chat that grounds every reply in the user's actual scorecard, so the advice is specific to them rather than generic.",
)
add_para(
    doc,
    "The product has two modes. SOLO is for a single user doing a self-assessment alone. ENTERPRISE is for a company admin who invites employees by email and controls which assessment areas each employee can answer. The KPIs are mapped to NIST CSF 2.0 and ISO 27001 control families, so the result is a real signal rather than a toy score.",
)
add_para(
    doc,
    "I built this project over five weeks as the sole developer. It is now fully deployed and works end to end. The Android APK runs standalone on a real phone, the API runs on Vercel in the Singapore region, the database is on Neon Postgres (also Singapore), the cache is on Upstash Redis (also Singapore), and real password-reset emails arrive in real inboxes through Brevo. The total cost to run the deployment is zero rupees per month because every service is on a free tier.",
)


# ----- 2. Problem -----
add_heading(doc, "2. Problem Statement")
add_para(
    doc,
    "Most small and mid-size companies do not actually know how good their cybersecurity is. They hear about ransomware in the news, they hear about data breaches, and they feel that they should do something. But they do not have a CISO. They do not have a security consultant. They do not even have a checklist short enough to actually use.",
)
add_para(
    doc,
    "The result is that they either ignore the problem entirely, or they panic buy security tools without knowing what they actually need. Both paths waste money and time, and neither path actually makes them safer.",
)
add_para(
    doc,
    "CyberScore solves this by giving the company a structured self-assessment that takes about thirty minutes. At the end the company knows three things: what is their current score, where are the weakest areas, and what are the next three actions to take. It is not a substitute for a real audit, but it is a cheap and useful first step that any founder or IT lead can complete on their phone in one sitting.",
)


# ----- 3. Solution -----
add_heading(doc, "3. Solution Overview")
add_para(doc, "CyberScore has three parts that work together.")
add_para(
    doc,
    "The mobile app is built with React Native and Expo. It is what the user actually interacts with. They sign in, take the assessment, see the scorecard, and chat with the advisor. Everything important happens here.",
)
add_para(
    doc,
    "The API is a Next.js server that runs the authentication, stores the answers, computes the scores, and streams the chat replies. This is the brain of the product.",
)
add_para(
    doc,
    "The database is a Postgres instance on Neon that holds the KPI catalogue, the user answers, the historical snapshots, and the chat history.",
)

add_heading(doc, "Key product features", level=2)
add_bullets(
    doc,
    [
        "Self-assessment across 46 KPIs in three categories (People, Process, Company).",
        "A live scorecard with the overall score, a per-category breakdown, and Red, Amber, Green bands per metric.",
        "Personalised remediation suggestions for each weak KPI.",
        "An AI advisor chat that knows the user's actual scorecard and gives industry-specific advice. Eight industries are covered: Banking, Healthcare, Technology, Manufacturing, Retail, Education, Government, and Other.",
        "A trend chart showing how the score has changed over time.",
        "Two user modes: SOLO for individuals and ENTERPRISE for admins who manage a team.",
        "Per-employee access control. The admin can give an employee access to only People questions, or any combination of the three levels.",
        "Optional two-factor authentication using any standard authenticator app.",
        "Shareable read-only scorecard URLs so the user can show their score to a CISO or insurance provider without giving them an account.",
        "A forgot-password flow that emails a six-digit code through Brevo.",
    ],
)


# ----- 4. Architecture -----
add_heading(doc, "4. System Architecture")
add_para(
    doc,
    "The system is a modular monolith. There is one API process, but the code inside it is split into clear modules: auth, scorecard, AI advisor, evidence, team, and so on. Each module exposes a flat surface and does not reach into another module's internals.",
)
add_para(
    doc,
    "I chose this shape over microservices because a single developer cannot maintain a microservices fleet for a five-week capstone. With one process, one deployment, and one log stream, the project stays simple. If the user base ever grows past about fifty thousand monthly active users, splitting modules into separate services would be straightforward because the boundaries are already clean.",
)

add_heading(doc, "Components in plain language", level=2)
add_para(
    doc,
    "The mobile app talks to the API over HTTPS with a JWT access token. The API reads and writes to Postgres for everything that needs to persist. The API reads and writes to Redis for rate limits and the KPI catalogue cache. The API talks to either the local rule-based advisor or the Anthropic Claude API for chat. The API talks to Brevo over HTTPS for sending one-time passwords and invites. A small external cron job pings the API health endpoint every two minutes to keep it warm.",
)

add_heading(doc, "What happens during one request", level=2)
add_para(doc, "When the user submits an assessment answer:")
add_bullets(
    doc,
    [
        "The mobile app sends a POST to /api/v1/kpis/submit with the KPI id and the answer.",
        "The API verifies the JWT and sets the database session variable app.current_org_id. This activates Row Level Security so the request can only see and write rows that belong to this organisation.",
        "The API writes the response row to the database.",
        "The API recomputes the scorecard live. The calculation is a pure function and runs in a few milliseconds.",
        "The API returns the updated scorecard JSON to the mobile app.",
        "The mobile app updates its React Query cache and shows the new score on the dashboard.",
    ],
)

add_heading(doc, "Tenant isolation through Row Level Security", level=2)
add_para(
    doc,
    "Every table that holds user data has RLS policies in Postgres. The policy says that a row is visible only when its orgId column matches the current session's app.current_org_id. If the API ever forgets to set that variable, the query would return zero rows instead of leaking another organisation's data. This is a safety net in case of a programming mistake.",
)


# ----- 5. Tech stack -----
add_heading(doc, "5. Technology Stack")
add_para(doc, "The major choices, with the reason for each.")

add_table(
    doc,
    headers=["Layer", "Choice", "Why"],
    rows=[
        ["Mobile UI", "Expo SDK 52, React Native 0.76, NativeWind 4", "One codebase for both iOS and Android. Native APIs available through Expo without writing native code."],
        ["Mobile state", "TanStack Query 5, Zustand 5", "Query for server cache and request lifecycle. Zustand for the auth store."],
        ["Server", "Next.js 15.5 (App Router)", "Route handlers without a separate Express layer. Excellent Vercel support."],
        ["Database", "PostgreSQL 16 with Prisma 6", "Relational data, schema as the source of truth, Row Level Security for tenant isolation."],
        ["Cache", "Redis on Upstash", "Rate-limit counters, KPI catalogue cache, sliding-window throttles."],
        ["Auth", "Argon2id plus JWT plus refresh tokens", "OWASP-recommended hashing, stateless access, revocable refresh."],
        ["AI advisor", "Local rule-based advisor by default. Anthropic Claude optional.", "Default is free and demoable. Anthropic available when budget allows."],
        ["Email", "Brevo HTTP API (free, 300 per day)", "HTTP over SMTP because Vercel serverless throttles raw TCP."],
        ["Validation", "Zod 3", "Single source of truth for request and response shapes shared between API and mobile."],
        ["Logging", "Pino", "Structured JSON logs, very low overhead."],
        ["Build", "Turbo 2 plus npm workspaces", "Cacheable lint, typecheck, and build across the monorepo."],
        ["Hosting", "Vercel + Neon + Upstash + Brevo + Expo EAS + cron-job.org", "All free tier. Zero rupees per month."],
    ],
    col_widths_inches=[1.3, 2.2, 3.0],
)


# ----- 6. Features in detail -----
add_heading(doc, "6. Features in Detail")

add_heading(doc, "Authentication and accounts", level=2)
add_bullets(
    doc,
    [
        "Three registration modes: SOLO, ENTERPRISE_ADMIN, and ENTERPRISE_EMPLOYEE.",
        "Passwords hashed with Argon2id. Every new password is checked against the Have I Been Pwned breach database to block known-compromised credentials.",
        "Email OTP verification on first registration. The code expires in ten minutes.",
        "Login issues a JWT access token (15 minute lifetime) and an opaque refresh token (7 day lifetime). Refresh tokens are bcrypt-hashed at rest and rotated on every use.",
        "Optional TOTP two-factor authentication. The secret is encrypted with AES-GCM before storage.",
        "Password reset through a six-digit OTP sent to the user's email through Brevo.",
    ],
)

add_heading(doc, "Assessment and scoring", level=2)
add_bullets(
    doc,
    [
        "46 KPIs across three levels: People, Process, Company.",
        "Each KPI has between three and five scoring tiers with concrete language so the user knows which one matches them.",
        "Scores are zero to one hundred per metric. Aggregates use weighted averages.",
        "Red is below 40, Amber is 40 to 70, Green is above 70.",
        "Historical snapshots are written on every recompute, so the trend chart shows real progress over time.",
    ],
)

add_heading(doc, "Enterprise mode", level=2)
add_bullets(
    doc,
    [
        "Admin can invite an employee by email. The invite carries the role and the set of allowed levels.",
        "Employee accepts the invite, sets their password, and joins the organisation.",
        "Admin can change an employee's allowed levels at any time. Useful when someone changes role.",
        "Admin sees a team dashboard with per-member completion percentage, score, and last activity timestamp.",
        "Aggregated team scorecard rolls up org-scope KPIs. If three employees give different answers to the same shared question, the admin sees the disagreement and the admin's answer wins for the rollup.",
        "Soft delete for revoked employees keeps the historical answer trail intact.",
        "Audit log records every team-related change with actor, timestamp, and before-and-after state.",
    ],
)

add_heading(doc, "AI advisor", level=2)
add_bullets(
    doc,
    [
        "Default is a local rule-based advisor that runs entirely inside the API. Zero external cost.",
        "Detects user intent from keywords: weakest, first priority, explain level, sector threats, sector controls, and others.",
        "Replies are grounded in the user's actual scorecard. The advice mentions specific KPIs by name.",
        "Sector knowledge primer for eight industries baked in: Banking (PCI DSS, RBI), Healthcare (HIPAA, ransomware), Technology (supply chain, SBOM), Manufacturing (OT segmentation), Retail (PCI DSS, e-commerce), Education (FERPA, student data), Government, and Other.",
        "Streamed word by word over Server-Sent Events with a 40ms delay between tokens. Looks and feels like a real LLM response.",
        "Switching to Anthropic Claude is a single environment variable flip. The same code path then streams from Sonnet 4.6 with prompt caching enabled and a daily budget guard.",
        "Per-user rate limit of 20 messages per minute via Redis sliding window.",
    ],
)


# ----- 7. Development process -----
add_heading(doc, "7. Development Process")
add_para(
    doc,
    "The five-week build was organised into five rough sprints, one per week. There was no formal Scrum or Jira because the team was one person. The sprints below are retrospective groupings of what got shipped each week.",
)

add_table(
    doc,
    headers=["Sprint", "Window", "Theme", "Outcome"],
    rows=[
        ["1", "15 to 21 Apr", "Bootstrap and auth", "Monorepo set up. Postgres schema with 21 tables. Argon2id auth with JWT plus refresh tokens. Mobile shell with onboarding, login, register, and verify-OTP screens."],
        ["2", "22 to 28 Apr", "KPI catalogue and scoring", "Extracted 46 KPIs from the source XLSX into the database. Built the pure-function scoring engine. Live dashboard with score ring. End-to-end assessment flow."],
        ["3", "29 Apr to 5 May", "Enterprise mode", "Per-user level permissions. Invite and accept-invite flow. Team admin screen with per-member score and edit-levels controls."],
        ["4", "6 to 12 May", "AI chat and password reset", "Anthropic streaming with prompt caching. Local rule-based advisor as the zero-cost default. Forgot-password flow with OTP email."],
        ["5", "13 to 19 May", "Deployment and demo", "Vercel API live in Singapore region. Neon Postgres in Singapore. Upstash Redis in Singapore. Brevo email. Android APK built and installed. Cold-start mitigation. Demo recording. Post-demo polish."],
    ],
    col_widths_inches=[0.5, 1.0, 1.6, 3.4],
)
add_para(doc, "Estimated active development hours: about 225 hours total, or roughly 45 hours per week.")


# ----- 8. Deployment -----
add_heading(doc, "8. Deployment")
add_para(
    doc,
    "The project is fully deployed and reachable from any internet-connected device. The total monthly cost is zero rupees because every service is on a free tier.",
)

add_table(
    doc,
    headers=["Component", "Service", "Region", "Free tier limit"],
    rows=[
        ["API", "Vercel (Hobby)", "Singapore (sin1)", "100 GB bandwidth"],
        ["Postgres", "Neon (Free)", "Singapore", "3 GB storage"],
        ["Redis", "Upstash (Free)", "Singapore", "10000 commands per day"],
        ["Email", "Brevo (Free)", "Global", "300 emails per day"],
        ["Mobile builds", "Expo EAS (Free)", "Cloud", "30 builds per month"],
        ["Keep-warm cron", "cron-job.org (Free)", "Cloud", "Unlimited jobs"],
    ],
    col_widths_inches=[1.4, 1.6, 1.4, 2.1],
)

add_heading(doc, "Important deployment decisions", level=2)
add_bullets(
    doc,
    [
        "Email is sent over HTTPS, not SMTP. Vercel serverless functions sometimes hang on outbound TCP for ports 465 and 587. The API auto-detects the email provider from the SMTP_PASS prefix and routes through the provider's REST API instead.",
        "Emails are sent inline, not queued. Originally the email worker was a separate process draining BullMQ jobs from Redis. On Vercel there is no worker process. The queue helpers were patched to call sendEmail directly.",
        "An external cron keeps the API warm. Vercel Hobby tier shuts down idle serverless functions after about five minutes. The first request after idle pays a 5 to 10 second cold-start tax. A free cron at cron-job.org pings the keepwarm endpoint every two minutes.",
        "Function region pinned to Singapore. Vercel originally placed our functions in a US region. Every Prisma query against Neon in Singapore paid a 200 to 400 millisecond transatlantic round trip. Pinning to sin1 brought per-query latency down to under 10 milliseconds.",
    ],
)
add_heading(doc, "Updates", level=2)
add_para(
    doc,
    "Every push to the main branch on GitHub triggers a Vercel auto-deploy. Build takes about three minutes. Database migrations are run manually from my laptop with prisma migrate deploy so that no preview deployment can touch the production schema by accident.",
)


# ----- 9. Testing -----
add_heading(doc, "9. Testing Approach")
add_para(
    doc,
    "Honest framing: there is no automated test suite in this version. Vitest is wired up in the API workspace but the test directories are empty. This is the single largest gap in the project. It is documented in the known issues file and called out as the first priority for the next batch.",
)
add_para(doc, "What is in place instead:")
add_bullets(
    doc,
    [
        "A scripted manual smoke test, about five minutes long, covering the critical paths: onboarding, login, dashboard, assessment, scorecard, AI chat, password reset.",
        "A scripted manual full regression, about thirty minutes long, covering all major flows.",
        "Live curl-based health checks. Direct API tests against the production endpoint with expected JSON shapes.",
        "Strict TypeScript and Zod schemas. The build catches a huge class of bugs before runtime.",
    ],
)
add_para(doc, "Priority test targets for the next batch when they start writing tests:")
add_bullets(
    doc,
    [
        "lib/scoring.ts. Pure functions, easy wins, high value.",
        "lib/scorecard.ts. Pure aggregation logic.",
        "lib/access.ts. Small but security-critical.",
        "lib/auth.ts. Token issue and revoke.",
        "Route handlers as integration tests with a Postgres test database.",
    ],
)


# ----- 10. Challenges -----
add_heading(doc, "10. Challenges and Lessons Learned")

add_heading(doc, "Challenge 1: The five-stage Vercel deployment", level=2)
add_para(
    doc,
    "What I expected to be a thirty-minute deploy turned into four hours of chasing five different problems in sequence.",
)
add_bullets(
    doc,
    [
        "Next.js 15.1.3 with React 19 had a known issue rendering the 404 page. Fix: added explicit app/not-found.tsx and pages/_error.tsx files.",
        "Vercel blocked the deploy because my git commit author email did not match the Vercel account email. Fix: rewrote git history with git filter-branch and force pushed.",
        "TypeScript ran on the scripts/ directory and tripped on a loose bcrypt import. Fix: added scripts/ to the tsconfig exclude list.",
        "Type errors only showed up on Vercel's clean install because of different module resolution. Fix: turned off Next's built-in TS and ESLint checks (we already run them in Turbo).",
        "Vercel security blocked the deploy with CVE-2025-66478 in next 15.1.3. Fix: upgraded next to 15.5.18.",
    ],
)
add_para(doc, "Lesson: deploy to the target on day two of the project, not week five. Cold-debugging a stack you have never deployed against is painful and slow.")

add_heading(doc, "Challenge 2: macOS Gatekeeper blocked Android Studio", level=2)
add_para(
    doc,
    "I tried to install Android Studio to build the APK locally. macOS Gatekeeper refused to open the app and said it was damaged. I spent ninety minutes trying various xattr commands and System Settings tricks. None of them worked.",
)
add_para(
    doc,
    "Then I realised I did not need Android Studio at all. Expo's EAS Build does the entire Android build in the cloud. I deleted Android Studio, set up EAS, and had a working APK twenty minutes later.",
)
add_para(doc, "Lesson: before going deep on a fix, ask yourself if you even need this thing. Knowing what not to do is half the battle.")

add_heading(doc, "Challenge 3: Email sending was harder than expected", level=2)
add_para(
    doc,
    "First attempt used Resend. Worked perfectly in curl tests. Then I read the fine print: Resend's free tier only delivers to the account owner's own email. Useless for a multi-user demo. Pivoted to Brevo (300 emails per day free, no recipient restriction). Then hit another problem: Vercel was throttling raw SMTP on ports 465 and 587. Switched to the HTTP API and everything started working.",
)
add_para(doc, "Lesson: read the free-tier restrictions of every SaaS before you integrate. Free does not always mean useful. And for serverless, always pick HTTP over SMTP.")

add_heading(doc, "Challenge 4: Cold-start lag killed the demo experience", level=2)
add_para(
    doc,
    "After the API was deployed and the APK was installed, every screen transition felt like the app was hanging for five to fifteen seconds. Root cause: Vercel Hobby tier shuts down idle serverless functions after about five minutes. Fix: a free external cron at cron-job.org pings the API every two minutes. Cost: zero rupees. Effect: the app feels instant.",
)
add_para(doc, "Lesson: serverless cold starts are real and they bite hard on free tiers. Always have a keep-warm strategy from day one.")

add_heading(doc, "Other lessons that would save the next batch time", level=2)
add_bullets(
    doc,
    [
        "Strict TypeScript from day one. The first week is annoying. After that, every refactor catches dozens of bugs at compile time.",
        "Single source of truth for schemas. Putting Zod schemas in a shared workspace package saved hours of request-shape drift debugging.",
        "Dev-mode escape hatches are gold. Returning OTPs inline in API responses when NODE_ENV is development meant local development did not depend on a working SMTP server.",
        "Test perceived latency from a real phone, not just API response time. My curl tests all returned in 200 milliseconds. The cold-start lag only showed up when I actually used the app on a phone.",
    ],
)


# ----- 11. Limitations -----
add_heading(doc, "11. Known Limitations")
add_para(doc, "In rough order of how much they matter:")
add_bullets(
    doc,
    [
        "No automated tests. Vitest is wired up but the test directories are empty. Manual testing covers smoke and regression flows. This is the single biggest gap.",
        "No CI pipeline. Lint and typecheck run only when I remember to run them locally. Vercel build catches type errors but does not run the full Turbo pipeline.",
        "Vercel cold starts are mitigated, not eliminated. The keep-warm cron stops most of them.",
        "Evidence file uploads not wired. The schema and routes exist but the R2 credentials are placeholders.",
        "Push notifications are dormant. The schema, the worker, and the SDK plumbing all exist, but no admin UI triggers them.",
        "Scorecard PDF export not built.",
        "Brevo's free tier rewrites the sender email to a generic relay domain. Display name CyberScore is preserved.",
        "No staging environment. One production database and one production deployment.",
    ],
)


# ----- 12. Future work -----
add_heading(doc, "12. Future Work")
add_para(
    doc,
    "What the next batch should pick up, in rough priority order. Detailed roadmap is in docs/roadmap.md.",
)

add_heading(doc, "Highest priority", level=2)
add_bullets(
    doc,
    [
        "Write tests for lib/scoring.ts, lib/scorecard.ts, lib/access.ts, lib/auth.ts. Target sixty percent coverage minimum.",
        "Wire up GitHub Actions CI to run lint, typecheck, and tests on every pull request.",
        "Set up Sentry for error tracking.",
        "Verify all RLS policies under realistic multi-tenant load.",
        "Set up a staging environment on a separate Neon branch.",
    ],
)

add_heading(doc, "Near-term features", level=2)
add_bullets(
    doc,
    [
        "Push notifications wired end to end.",
        "Scorecard PDF export through the email worker.",
        "Multi-framework view toggle on the scorecard.",
        "Industry benchmarks surfaced on the dashboard.",
        "Per-employee snapshot history so employees see their own trend lines.",
        "CSV export of the team scorecard.",
    ],
)

add_heading(doc, "Bigger items", level=2)
add_bullets(
    doc,
    [
        "Single Sign On (OIDC plus SAML) for enterprise customers.",
        "Continuous re-assessment (prompt users to refresh stale answers periodically).",
        "Compliance report templates (filled SOC 2 or ISO 27001 checklists from the scorecard).",
        "Webhook subscriptions for events.",
        "A read-only web dashboard.",
    ],
)


# ----- 13. Conclusion -----
add_heading(doc, "13. Conclusion")
add_para(
    doc,
    "CyberScore started as a five-week capstone and ended as a real, working, deployed product. The mobile app installs on a real Android phone. The API runs on Vercel. The database lives in Singapore. Emails arrive in real inboxes. The full happy path works end to end.",
)
add_para(
    doc,
    "What I am most proud of is that the architecture is clean enough that a new developer can pick it up without three weeks of onboarding. The schemas are shared between API and mobile. The scoring logic lives in pure functions. The auth is built on standard primitives. Row Level Security is in place. The local advisor was a thoughtful response to a budget constraint rather than a hack.",
)
add_para(
    doc,
    "What I am most aware of is that there are no automated tests. That is the next batch's first task and I would do it myself if I had another week.",
)
add_para(
    doc,
    "Thanks to Mohan Ram C for the mentorship, especially the conversation about the chat advisor that led to the sector knowledge primer design. The product is better for it.",
)


# ----- 14. References -----
add_heading(doc, "14. References and Links")
add_table(
    doc,
    headers=["Reference", "Where"],
    rows=[
        ["Live API", "https://cyberscore-api.vercel.app"],
        ["Health check", "https://cyberscore-api.vercel.app/api/v1/health"],
        ["GitHub repository", "https://github.com/saanvivishal/cyberscore"],
        ["Demo video", "IIIT Bangalore SharePoint (link in handover email)"],
        ["Demo login email", "saanvi.vishal@iiitb.ac.in"],
        ["Demo login password", "cyberscore-demo-2026"],
        ["Handover checklist", "HANDOVER_CHECKLIST.md"],
        ["Architecture document", "docs/architecture.md"],
        ["System design document", "docs/system-design.md"],
        ["Requirements document", "docs/requirements.md"],
        ["Database document", "docs/database.md"],
        ["Deployment runbook", "docs/deployment.md"],
        ["Known issues", "docs/known-issues.md"],
        ["Future roadmap", "docs/roadmap.md"],
        ["Sprint history", "docs/sprints.md"],
        ["Access and credentials", "docs/access-credentials.md"],
        ["Handover notes", "docs/handover-notes.md"],
        ["Manual test plan", "docs/testing-manual.md"],
        ["Changelog", "CHANGELOG.md"],
        ["Contributing guide", "CONTRIBUTING.md"],
    ],
    col_widths_inches=[2.2, 4.5],
)


# ============================================================
# Part 2: Slides outline
# ============================================================
doc.add_page_break()
add_heading(doc, "Part Two: Presentation Slides Outline", level=0)
add_para(
    doc,
    "Six slides for a 10 to 15 minute talk.",
    size=14,
    italic=True,
    color=MUTED,
    align="center",
)
doc.add_paragraph()
add_para(
    doc,
    "Each slide block below shows the title, what goes on the slide, and a short speaker note for what to say while the slide is up. Keep slide text minimal so the audience listens, not reads. The demo video on slide four is the centrepiece of the talk.",
)


# Slide 1
add_heading(doc, "Slide 1: Title")
add_para(doc, "Title: CyberScore. A Cybersecurity Health Scorecard for Organisations.", bold=True)
add_para(doc, "On the slide:", bold=True)
add_bullets(
    doc,
    [
        "Project name in large type: CyberScore",
        "Tagline: Self-assessment cybersecurity in your pocket",
        "Your name and roll number: Saanvi Vishal, IMT2021043",
        "Mentor: Mohan Ram C, FISST",
        "Institution: International Institute of Information Technology, Bangalore",
        "Date of presentation",
    ],
)
add_para(doc, "Speaker note (about 30 seconds):", bold=True)
add_para(
    doc,
    "Introduce yourself, name the project, name the mentor and the institution. One sentence about what CyberScore is. A mobile app that lets a company check its cybersecurity health in thirty minutes and get specific advice on what to fix next.",
    italic=True,
    color=MUTED,
)


# Slide 2
add_heading(doc, "Slide 2: The Problem and the Solution")
add_para(doc, "Title: The Problem, and What CyberScore Does", bold=True)
add_para(doc, "On the slide (left side, the problem):", bold=True)
add_bullets(
    doc,
    [
        "Small and mid-size companies do not know how good their cybersecurity is.",
        "They have no CISO, no consultant, no usable checklist.",
        "Result: either ignore the problem or panic-buy security tools.",
    ],
)
add_para(doc, "On the slide (right side, the solution):", bold=True)
add_bullets(
    doc,
    [
        "Mobile app. Thirty-minute self-assessment. 46 questions.",
        "Score across three areas: People, Process, Company.",
        "AI advisor that knows your actual score and your industry.",
        "Two modes: SOLO for one user, ENTERPRISE for an admin with employees.",
    ],
)
add_para(doc, "Speaker note (about 90 seconds):", bold=True)
add_para(
    doc,
    "Describe the problem in plain language. A founder running a small fintech with five engineers cannot afford a security audit. They know they should do something but they do not know what. CyberScore is a structured self-assessment that gives them a score, shows them where they are weak, and tells them the next three actions. Not a replacement for a real audit, but a cheap and useful first step.",
    italic=True,
    color=MUTED,
)


# Slide 3
add_heading(doc, "Slide 3: How It Is Built")
add_para(doc, "Title: How It Is Built", bold=True)
add_para(doc, "On the slide (top half, the architecture diagram):", bold=True)
add_para(
    doc,
    "Use the system context diagram from docs/architecture.md. Mobile app on the left, API in the middle, Postgres and Redis on the right, plus arrows to Brevo for email and the local advisor or Anthropic for AI.",
)
add_para(doc, "On the slide (bottom half, tech stack list):", bold=True)
add_bullets(
    doc,
    [
        "Mobile: React Native with Expo SDK 52",
        "API: Next.js 15.5 on Vercel",
        "Database: PostgreSQL 16 with Prisma. Row Level Security for tenant isolation.",
        "Cache: Redis on Upstash",
        "AI: local rule-based advisor (default), Anthropic Claude (optional)",
        "Auth: Argon2id passwords, JWT, refresh tokens",
        "Shared types: Zod schemas in a workspace package",
    ],
)
add_para(doc, "Speaker note (about 90 seconds):", bold=True)
add_para(
    doc,
    "Walk the audience through the diagram once. Then highlight two design choices: one, it is a modular monolith because a single developer cannot maintain microservices; two, every request and response shape is shared between API and mobile through a single Zod schema package so there is no drift between the front end and the back end.",
    italic=True,
    color=MUTED,
)


# Slide 4
add_heading(doc, "Slide 4: Live Demo")
add_para(doc, "Title: Live Demo", bold=True)
add_para(doc, "On the slide:", bold=True)
add_bullets(
    doc,
    [
        "One or two big screenshots: the dashboard with the score ring, and the AI chat with a sector-specific reply",
        "A QR code or short URL pointing to the demo video",
        "A short feature reminder list (six bullets max)",
    ],
)
add_para(doc, "Speaker note (about four to five minutes, while playing the demo):", bold=True)
add_para(
    doc,
    "Play the demo video. Talk over the top of it for any parts that need explanation. Highlight three things during playback. One: the onboarding gives a clean first impression. Two: every action updates the scorecard live. Three: the AI advisor mentions specific KPIs by name because it has the real scorecard data in its context.",
    italic=True,
    color=MUTED,
)


# Slide 5
add_heading(doc, "Slide 5: Deployment, Process, and Cost")
add_para(doc, "Title: Live in Production at Zero Cost", bold=True)
add_para(doc, "On the slide (left column, what is deployed):", bold=True)
add_bullets(
    doc,
    [
        "API: Vercel Hobby tier, Singapore region (free)",
        "Database: Neon Postgres in Singapore (free, 3 GB)",
        "Cache: Upstash Redis in Singapore (free, 10k commands per day)",
        "Email: Brevo HTTP API (free, 300 per day)",
        "Mobile: APK built and distributed through Expo EAS",
        "Keep-warm: external cron pings every two minutes (free)",
        "Total monthly cost: zero rupees",
    ],
)
add_para(doc, "On the slide (right column, how I built it):", bold=True)
add_bullets(
    doc,
    [
        "Five weekly sprints, single developer",
        "Sprint 1: bootstrap and auth",
        "Sprint 2: KPI catalogue and scoring",
        "Sprint 3: enterprise mode and team management",
        "Sprint 4: AI chat (Anthropic plus local advisor)",
        "Sprint 5: production deployment and demo",
        "About 225 active development hours total",
    ],
)
add_para(doc, "Speaker note (about 90 seconds):", bold=True)
add_para(
    doc,
    "The deployment story is something I am proud of because it works end to end on free tiers. Mention the cold-start mitigation: Vercel idles functions after five minutes, so I set up a free external cron to ping the API every two minutes. This keeps the app feeling fast without paying for Vercel Pro. Then walk through the five sprints in one line each.",
    italic=True,
    color=MUTED,
)


# Slide 6
add_heading(doc, "Slide 6: Lessons Learned and Thanks")
add_para(doc, "Title: What I Learned and What Comes Next", bold=True)
add_para(doc, "On the slide (lessons, four bullets):", bold=True)
add_bullets(
    doc,
    [
        "Deploy on day two, not week five. Cold-debugging a stack you have never deployed against is painful.",
        "Read the free-tier fine print before you integrate. Free does not always mean useful.",
        "For serverless, pick HTTP APIs over SMTP. Outbound TCP is unreliable.",
        "Strict TypeScript plus shared Zod schemas paid back the setup cost in week one.",
    ],
)
add_para(doc, "On the slide (what is left):", bold=True)
add_bullets(
    doc,
    [
        "Write tests (Vitest is wired, directories are empty)",
        "Wire GitHub Actions CI",
        "Push notifications and scorecard PDF export",
        "Sentry for error tracking",
    ],
)
add_para(doc, "On the slide (thanks):", bold=True)
add_bullets(
    doc,
    [
        "Mohan Ram C for the mentorship and the sector primer idea that shaped the AI advisor design",
        "IIIT Bangalore for the platform",
        "Questions and answers",
    ],
)
add_para(doc, "Speaker note (about 90 seconds):", bold=True)
add_para(
    doc,
    "Be honest about what is not done. No tests. No CI. Both go in the roadmap. Then close with a real thank you to your mentor and offer to take questions. If the audience is engaged, this is where the conversation gets interesting.",
    italic=True,
    color=MUTED,
)


# ----- Tips at the end -----
doc.add_page_break()
add_heading(doc, "Tips for Delivering the Presentation")
add_bullets(
    doc,
    [
        "Keep slide text short. The audience should listen to you, not read at you.",
        "Spend most of the time on slide 4 (the demo). The video is the strongest part of the project.",
        "Practice running the demo video at least twice before the real presentation. Know where to pause and add commentary.",
        "If the demo video link does not load on the day, have the APK installed on your own phone as a backup and project that screen.",
        "Be honest about the gaps (no tests, no CI). Supervisors respect honesty more than false polish.",
        "Have the GitHub link and the handover checklist ready in a browser tab in case anyone asks for code or documentation during questions.",
    ],
)
add_para(doc, "")
add_para(doc, "End of document. Saanvi Vishal, 19 May 2026.", italic=True, color=MUTED, align="center")


# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
